"""Exporta informe-tecnico-final.md a DOCX y PDF con formato academico."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS = Path(__file__).resolve().parent.parent
MD_PATH = DOCS / "informe-tecnico-final.md"
OUT_DOCX = DOCS / "informe-tecnico-final.docx"
OUT_PDF = DOCS / "informe-tecnico-final.pdf"

FONT = "Times New Roman"
BODY_SIZE = Pt(12)
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)


def read_md() -> str:
    return MD_PATH.read_text(encoding="utf-8")


def strip_latex(text: str) -> str:
    text = re.sub(r"\\\[(.*?)\\\]", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\\\((.*?)\\\)", r"\1", text)
    text = text.replace("*", "")
    return text.strip()


def shade_paragraph(paragraph, fill: str = "F2F4F7") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12), ("Heading 4", 12)]:
        style = doc.styles[name]
        style.font.name = FONT
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = HEADING_COLOR
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_centered(doc: Document, text: str, size: int, bold: bool = False, space_after: int = 8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = HEADING_COLOR
    p.paragraph_format.space_after = Pt(space_after)


def build_cover(doc: Document) -> None:
    add_centered(doc, "UNIVERSIDAD ESAN", 13, bold=True, space_after=4)
    add_centered(doc, "Facultad de Ciencias de la Computación", 11, space_after=18)
    add_centered(doc, "INFORME TÉCNICO FINAL", 16, bold=True, space_after=10)
    add_centered(
        doc,
        "Estudio, implementación, integración con base de datos\ny simulación visual del Scapegoat Tree",
        12,
        space_after=24,
    )

    meta = [
        ("Curso", "Algoritmos y Estructura de Datos (2026-1)"),
        ("Docente", "Marks Calderón Niquin"),
        ("Integrantes", "[Nombre Apellido — Código] · [Nombre Apellido — Código] · [Nombre Apellido — Código]"),
        ("Repositorio", "https://github.com/Fravely/Spacegoat-tree"),
        ("Base de datos", "Microsoft SQL Server (go-mssqldb)"),
        ("Dataset", "dbo.Productos — inventario de productos tecnológicos"),
        ("Implementación", "Go 1.23 · Vue.js 3"),
        ("Fecha", "Junio 2026"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.autofit = True
    for i, (label, value) in enumerate(meta):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for c in range(2):
            for p in row.cells[c].paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(11)
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], "E8EDF5")

    doc.add_paragraph()
    p = doc.add_paragraph("Trabajo Final — Semestre Académico 2026-1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = FONT
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True
    add_page_break(doc)


def parse_blocks(content: str) -> list[tuple[str, str]]:
    content = re.sub(r"<!--.*?-->", "", content, flags=re.DOTALL)
    blocks: list[tuple[str, str]] = []
    lines = content.splitlines()
    i = 0
    skip_until_section_one = True

    while i < len(lines):
        line = lines[i]

        if skip_until_section_one:
            if line.startswith("# 1."):
                skip_until_section_one = False
            else:
                i += 1
                continue

        if line.strip() == "<!-- PAGEBREAK -->":
            blocks.append(("pagebreak", ""))
            i += 1
            continue
        if line.startswith("```"):
            lang = line[3:].strip()
            i += 1
            body = []
            while i < len(lines) and not lines[i].startswith("```"):
                body.append(lines[i])
                i += 1
            blocks.append((f"code:{lang}", "\n".join(body)))
            i += 1
            continue
        if line.startswith("#"):
            blocks.append(("heading", line))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                table.append(lines[i])
                i += 1
            blocks.append(("table", "\n".join(table)))
            continue
        if line.strip().startswith("!["):
            blocks.append(("image", line.strip()))
            i += 1
            continue
        if line.strip() == "---":
            blocks.append(("separator", ""))
            i += 1
            continue
        if line.strip():
            para = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("|") and not lines[i].startswith("```") and lines[i].strip() not in ("---", "<!-- PAGEBREAK -->") and not lines[i].strip().startswith("!["):
                para.append(lines[i])
                i += 1
            blocks.append(("para", "\n".join(para)))
            continue
        i += 1
    return blocks


def add_table(doc: Document, raw: str) -> None:
    rows = [r.strip() for r in raw.splitlines() if r.strip()]
    data = [[c.strip() for c in row.strip("|").split("|")] for row in rows]
    data = [r for r in data if not all(set(c) <= set("-:") for c in r)]
    if len(data) < 1:
        return

    cols = len(data[0])
    table = doc.add_table(rows=len(data), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r, row in enumerate(data):
        for c, val in enumerate(row):
            if c >= cols:
                break
            cell = table.rows[r].cells[c]
            cell.text = strip_latex(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = FONT
                    run.font.size = Pt(10)
                    if r == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            if r == 0:
                set_cell_shading(cell, "1A1A2E")

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    doc.add_paragraph()


def add_code_block(doc: Document, lang: str, content: str) -> None:
    label = doc.add_paragraph()
    run = label.add_run(f"Código fuente ({lang or 'texto'})")
    run.font.name = FONT
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x47, 0x54, 0x67)

    for line in content.splitlines():
        p = doc.add_paragraph(line if line else " ")
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(0)
        shade_paragraph(p, "F4F6F8")
        if p.runs:
            r = p.runs[0]
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x10, 0x18, 0x28)
    doc.add_paragraph()


def add_mermaid_note(doc: Document, content: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Figura algorítmica (diagrama de flujo). ")
    run.font.name = FONT
    run.font.size = Pt(11)
    run.italic = True
    note = p.add_run(
        "La representación gráfica completa del diagrama está disponible en la versión Markdown del informe. "
        "A continuación se resume la lógica representada."
    )
    note.font.name = FONT
    note.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

    keywords = []
    for line in content.splitlines():
        if "-->" in line or "[" in line:
            keywords.append(line.strip()[:80])
    for line in keywords[:6]:
        bullet = doc.add_paragraph(line, style="List Bullet")
        for r in bullet.runs:
            r.font.name = FONT
            r.font.size = Pt(10)


def add_figure(doc: Document, caption: str, rel: str) -> None:
    img = DOCS / rel.replace("./", "")
    if not img.exists():
        return
    doc.add_picture(str(img), width=Inches_safe(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    for r in cap.runs:
        r.font.name = FONT
        r.font.size = Pt(10)
        r.italic = True
        r.font.color.rgb = RGBColor(0x47, 0x54, 0x67)


def Inches_safe(cm_val: float):
    from docx.shared import Inches
    return Inches(cm_val / 2.54)


def add_footer_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = FONT
        run.font.size = Pt(10)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)


def build_docx(blocks: list[tuple[str, str]]) -> None:
    doc = Document()
    configure_document(doc)
    build_cover(doc)

    # Resumen e indice desde el MD original (seccion pre-1)
    raw = read_md()
    resumen_match = re.search(r"## Resumen\n\n(.*?)\n\n\*\*Palabras clave", raw, re.DOTALL)
    if resumen_match:
        doc.add_heading("Resumen", level=1)
        p = doc.add_paragraph(strip_latex(resumen_match.group(1)))
        p.paragraph_format.first_line_indent = Cm(1.0)
        kw = doc.add_paragraph()
        r = kw.add_run("Palabras clave: ")
        r.bold = True
        r.font.name = FONT
        r2 = kw.add_run(
            "Scapegoat Tree, árbol binario de búsqueda, balance por reconstrucción, Go, SQL Server, "
            "complejidad amortizada, estructuras de datos."
        )
        r2.font.name = FONT
        add_page_break(doc)

    doc.add_heading("Índice general", level=1)
    items = [
        "Introducción y marco teórico",
        "Arquitectura del sistema y lógica del código (backend en Go)",
        "Integración con la base de datos real",
        "Diseño de la API REST",
        "Interfaz interactiva y simulación (frontend en Vue.js)",
        "Benchmarking y análisis de complejidad asintótica",
        "Gestión del proyecto y reporte de commits",
        "Conclusiones y recomendaciones",
        "Referencias y anexos",
    ]
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1.0)
        for r in p.runs:
            r.font.name = FONT
            r.font.size = Pt(12)
    add_page_break(doc)

    figure_no = 0
    table_no = 0

    for kind, content in blocks:
        if kind == "pagebreak":
            add_page_break(doc)
        elif kind == "heading":
            level = len(content) - len(content.lstrip("#"))
            text = strip_latex(content.lstrip("#").strip())
            doc.add_heading(text, level=min(level, 4))
        elif kind == "para":
            text = strip_latex(content)
            if text.startswith("Palabras clave"):
                continue
            p = doc.add_paragraph(text)
            if not text.startswith("-") and len(text) > 60:
                p.paragraph_format.first_line_indent = Cm(1.0)
        elif kind.startswith("code:"):
            lang = kind.split(":", 1)[1]
            if lang == "mermaid":
                figure_no += 1
                add_mermaid_note(doc, content)
            else:
                add_code_block(doc, lang, content)
        elif kind == "table":
            table_no += 1
            add_table(doc, content)
        elif kind == "image":
            m = re.search(r"!\[([^\]]*)\]\(([^)]+)\)", content)
            if m:
                figure_no += 1
                caption = m.group(1) or f"Figura {figure_no}"
                if not caption.lower().startswith("figura"):
                    caption = f"Figura {figure_no}. {caption}"
                add_figure(doc, caption, m.group(2))
        elif kind == "separator":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)

    add_footer_page_numbers(doc)
    doc.save(OUT_DOCX)


def main() -> None:
    blocks = parse_blocks(read_md())
    build_docx(blocks)
    print(f"DOCX: {OUT_DOCX}")
    try:
        from docx2pdf import convert
        convert(str(OUT_DOCX), str(OUT_PDF))
        print(f"PDF:  {OUT_PDF}")
    except Exception as exc:
        print(f"No se pudo generar PDF automaticamente: {exc}")
        print("Abra el DOCX en Word y exporte a PDF manualmente.")


if __name__ == "__main__":
    main()
